#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
MCP Docx Processing Service
Provides various operations for docx documents, including querying, adding, modifying, deleting, and font style settings
Implemented using the official MCP library
"""

import uvicorn # ADD THIS
import os
import tempfile
import logging
import traceback
import sys
from contextlib import asynccontextmanager
from typing import AsyncIterator, Dict, Any, Optional

from mcp.server.fastmcp import FastMCP, Context
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from typing import AsyncIterator, Dict, Any, Optional, List # Make sure List is here
from docx.text.run import Run as DocxRun # For type hinting, aliased to avoid confusion if you have a variable named Run

# --- Additions for Direct HTTP Endpoints ---
from starlette.applications import Starlette as StarletteApp # Alias to avoid confusion with mcp.sse_app() returning Starlette
from starlette.responses import JSONResponse
from starlette.routing import Route as HttpRoute, Mount
from starlette.requests import Request as StarletteRequest # Alias for clarity
# --- End Additions ---

# Configure logging with more detailed information
logging.basicConfig(
    level=logging.DEBUG,  # Changed to DEBUG for more detailed logs
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s - %(filename)s:%(lineno)d",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "docx_mcp_server.log")),
        logging.StreamHandler(sys.stderr)  # Changed to stdout for better visibility
    ]
)
logger = logging.getLogger("DocxMCPServer")

# Add debug logging for startup
logger.debug("Starting MCP Docx server...")
logger.debug(f"Python version: {sys.version}")
logger.debug(f"Current working directory: {os.getcwd()}")
logger.debug(f"Python path: {sys.path}")

# Create a state file for restoring state when MCP service restarts
CURRENT_DOC_FILE = os.path.join(tempfile.gettempdir(), "docx_mcp_current_doc.txt")

class DocxProcessor:
    """Class for processing Docx documents, implementing various document operations"""
    
    def __init__(self):
        self.documents: Dict[str, Document] = {}  # Store opened documents with type hint
        self.current_document: Optional[Document] = None # Type hinting
        self.current_file_path: Optional[str] = None # Type hinting
        
        # Try to load current document from state file
        self._load_current_document()
    
    def _load_current_document(self):
        """Load current document from state file"""
        if not os.path.exists(CURRENT_DOC_FILE):
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'r', encoding='utf-8') as f:
                file_path = f.read().strip()
            
            if file_path and os.path.exists(file_path):
                try:
                    self.current_file_path = file_path
                    self.current_document = Document(file_path)
                    self.documents[file_path] = self.current_document
                    return True
                except Exception as e:
                    logger.error(f"Failed to load document at {file_path}: {e}")
                    # Delete invalid state file to prevent future loading attempts
                    try:
                        os.remove(CURRENT_DOC_FILE)
                        logger.info(f"Removed invalid state file pointing to {file_path}")
                    except Exception as e_remove:
                        logger.error(f"Failed to remove state file: {e_remove}")
            else:
                # Delete invalid state file if path is empty or file doesn't exist
                try:
                    os.remove(CURRENT_DOC_FILE)
                    logger.info("Removed invalid state file with non-existent document path")
                except Exception as e_remove:
                    logger.error(f"Failed to remove state file: {e_remove}")
        except Exception as e:
            logger.error(f"Failed to load current document: {e}")
            # Delete corrupted state file
            try:
                os.remove(CURRENT_DOC_FILE)
                logger.info("Removed corrupted state file")
            except Exception as e_remove:
                logger.error(f"Failed to remove state file: {e_remove}")
        
        return False
    
    def _save_current_document_path_state(self): # Renamed for clarity
        """Save current document path to state file"""
        if not self.current_file_path:
            logger.debug("_save_current_document_path_state: No current_file_path to save.")
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'w', encoding='utf-8') as f:
                f.write(self.current_file_path)
            logger.debug(f"_save_current_document_path_state: Saved path '{self.current_file_path}' to state file.")
            return True
        except Exception as e:
            logger.error(f"Failed to save current document path state: {e}", exc_info=True)
        
        return False
    
    def save_state(self): # This method now correctly describes saving the DOCX and its path state
        """Save the current document to its file and update the state file with its path."""
        if self.current_document and self.current_file_path:
            try:
                self.current_document.save(self.current_file_path)
                logger.info(f"Document saved to: {self.current_file_path}")
                self._save_current_document_path_state() # Then save the path to the state file
            except Exception as e:
                logger.error(f"Failed to save current document or its state: {e}", exc_info=True)
        else:
            logger.info("save_state: No current document or file path, nothing to save.")
    
    def load_state(self):
        """Load processor state"""
        self._load_current_document()

    # ----- NEW METHODS FOR STRUCTURED CONTENT AND EDITING -----

    def get_structured_document_content_internal(self) -> List[Dict[str, Any]]:
        """
        Internal logic to extract structured content from the current document.
        Processes paragraphs, headings, and tables with their run-level formatting.
        Includes original indices for paragraphs and tables.
        """
        if not self.current_document:
            logger.warning("get_structured_document_content_internal: No active document.")
            raise ValueError("No active document to process.")
        
        doc: DocumentObject = self.current_document
        content_blocks = []
        block_id_counter = 0 

        # Pre-build maps from element ID to original index for efficient lookup
        para_element_map = {p._element: i for i, p in enumerate(doc.paragraphs)}
        table_element_map = {t._element: i for i, t in enumerate(doc.tables)}
        
        logger.debug(f"Starting extraction from document body. Found {len(doc.element.body)} direct children.")
        for child_element in doc.element.body:
            if isinstance(child_element, CT_P):
                # --- Process Paragraph --- 
                try:
                    para_object = Paragraph(child_element, doc)
                    # Look up original index using the element map
                    doc_paragraph_index = para_element_map.get(child_element)
                    logger.debug(f"Processing direct child: Paragraph (ID: p_block_{block_id_counter}, Orig Index: {doc_paragraph_index})")
                except Exception as e_para_map:
                    logger.warning(f"Could not map CT_P element back to Paragraph object: {e_para_map}. Skipping element.")
                    continue

                block_info: Dict[str, Any] = {
                    "id": f"p_block_{block_id_counter}",
                    "doc_paragraph_index": doc_paragraph_index,
                    "overall_block_index": block_id_counter, 
                    "type": "paragraph",
                    "text": para_object.text,
                    "style_name": para_object.style.name if para_object.style else "Normal",
                    "alignment": None, 
                    "runs": []
                }
                # DEBUGGING page_break_before for paragraphs
                pPr = para_object._element.pPr
                pbf_value = pPr.pageBreakBefore is not None if pPr is not None else False
                # --- NEW: Check for run-level page breaks (w:br w:type="page") ---
                page_break_in_run = False
                for run in para_object.runs:
                    for br in run._element.findall('.//w:br', namespaces=run._element.nsmap):
                        if br.get(qn('w:type')) == 'page':
                            page_break_in_run = True
                            break
                    if page_break_in_run:
                        break
                block_info["page_break_before"] = pbf_value or page_break_in_run

                if para_object.style and para_object.style.name.lower().startswith('heading'):
                    block_info["type"] = "heading"
                    # Extract heading level if possible
                    try:
                        level_str = para_object.style.name.split()[-1]
                        block_info["heading_level"] = int(level_str) if level_str.isdigit() else 0
                    except:
                        block_info["heading_level"] = 0 

                # Extract alignment
                if para_object.alignment is not None:
                    try:
                        alignment_map = {
                            WD_PARAGRAPH_ALIGNMENT.LEFT: "LEFT", WD_PARAGRAPH_ALIGNMENT.CENTER: "CENTER",
                            WD_PARAGRAPH_ALIGNMENT.RIGHT: "RIGHT", WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "JUSTIFY",
                            WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: "DISTRIBUTE", WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY: "THAI_JUSTIFY"
                        }
                        block_info["alignment"] = alignment_map.get(para_object.alignment)
                    except Exception as e_align:
                        logger.warning(f"Could not determine alignment for paragraph block {block_id_counter}: {e_align}")

                # Extract runs
                for run_element in para_object.runs:
                    font = run_element.font
                    block_info["runs"].append({
                        "text": run_element.text,
                        "bold": run_element.bold if run_element.bold is not None else False,
                        "italic": run_element.italic if run_element.italic is not None else False,
                        "underline": run_element.underline if run_element.underline is not None else False,
                        "font_name": font.name,
                        "font_size_pt": font.size.pt if font.size and hasattr(font.size, 'pt') else None,
                        "font_color_rgb": str(font.color.rgb) if font.color and font.color.rgb else None
                    })
                
                content_blocks.append(block_info)
                block_id_counter += 1
            
            elif isinstance(child_element, CT_Tbl):
                # --- Process Table --- 
                try:
                    table_object = Table(child_element, doc)
                    doc_table_index = table_element_map.get(child_element)
                    logger.debug(f"Processing direct child: Table (ID: t_meta_{block_id_counter}, Orig Index: {doc_table_index})")
                except Exception as e_tbl_map:
                    logger.warning(f"Could not map CT_Tbl element back to Table object: {e_tbl_map}. Skipping element.")
                    continue
                
                table_meta_block_info: Dict[str, Any] = {
                    "id": f"table_meta_{block_id_counter}",
                    "doc_table_index": doc_table_index,
                    "overall_block_index": block_id_counter, 
                    "type": "table_metadata",
                    "num_rows": len(table_object.rows),
                    "num_cols": len(table_object.columns), # This is logical columns from tblGrid
                    "style_name": table_object.style.name if table_object.style else "TableGrid",
                }
                content_blocks.append(table_meta_block_info)
                current_block_id_for_table_meta = block_id_counter # Save for logging cell association
                block_id_counter +=1

                # --- Advanced Table Cell Processing with Merge Handling ---
                # Determine actual number of columns from the first row if tblGrid is unreliable
                # This is a fallback and might not be perfect for all complex tables.
                actual_cols = len(table_object.columns) # Logical columns based on tblGrid
                if actual_cols == 0 and len(table_object.rows) > 0:
                    try:
                        actual_cols = len(table_object.rows[0].cells)
                        logger.debug(f"  Table {doc_table_index} has 0 logical columns from tblGrid, using actual cell count from first row: {actual_cols}")
                    except IndexError:
                        logger.warning(f"  Table {doc_table_index} has 0 logical columns and 0 rows. Cannot determine column count.")
                        actual_cols = 1 # Avoid division by zero later if table is truly empty
                elif actual_cols == 0:
                     logger.warning(f"  Table {doc_table_index} has 0 logical columns and 0 rows. Setting columns to 1.")
                     actual_cols = 1

                grid_cell_occupier = [[None for _ in range(actual_cols)] for _ in range(len(table_object.rows))]

                for r_idx in range(len(table_object.rows)):
                    for c_idx in range(actual_cols):
                        if grid_cell_occupier[r_idx][c_idx] is not None:
                            # This logical cell is already part of a processed merged cell
                            logger.debug(f"  Skipping grid cell ({r_idx},{c_idx}) for table {doc_table_index} - already occupied by {grid_cell_occupier[r_idx][c_idx]}")
                            continue

                        try:
                            current_cell_obj: _Cell = table_object.cell(r_idx, c_idx)
                        except IndexError:
                            logger.error(f"  IndexError accessing cell ({r_idx},{c_idx}) in table {doc_table_index}. Max rows: {len(table_object.rows)}, Max cols: {actual_cols}. Skipping this grid position.")
                            grid_cell_occupier[r_idx][c_idx] = 'INDEX_ERROR' # Mark to avoid reprocessing
                            continue

                        # Check for vertical merge continuation from cell properties
                        tcPr = current_cell_obj._tc.tcPr
                        v_merge_val = None
                        if tcPr is not None:
                            v_merge_elem = tcPr.vMerge
                            if v_merge_elem is not None:
                                v_merge_val = v_merge_elem.get(qn('w:val'))
                        
                        if v_merge_val is not None and v_merge_val != 'restart':
                            # This cell is a vertical continuation of a cell from a previous row.
                            # Its content should be part of the 'restart' cell. Mark as occupied by cell above.
                            # Find the restart cell by looking up in the grid_cell_occupier map.
                            occupier_above = None
                            if r_idx > 0 and grid_cell_occupier[r_idx-1][c_idx] is not None and grid_cell_occupier[r_idx-1][c_idx] != 'INDEX_ERROR':
                                occupier_above = grid_cell_occupier[r_idx-1][c_idx]
                            grid_cell_occupier[r_idx][c_idx] = occupier_above if occupier_above else (r_idx-1, c_idx) # Fallback if restart not found in map
                            logger.debug(f"  Skipping cell ({r_idx},{c_idx}) in table {doc_table_index} - vMerge continuation. Marked as occupied by {grid_cell_occupier[r_idx][c_idx]}.")
                            continue

                        # This is a primary cell (top-left of a visual block)
                        colspan = 1
                        if tcPr is not None and tcPr.gridSpan is not None and tcPr.gridSpan.val is not None:
                            try: colspan = int(tcPr.gridSpan.val) 
                            except ValueError: pass
                        
                        rowspan = 1
                        if v_merge_val == 'restart':
                            # Calculate actual rowspan by checking cells below in the same column
                            for rn_idx in range(r_idx + 1, len(table_object.rows)):
                                try:
                                    next_row_cell_obj: _Cell = table_object.cell(rn_idx, c_idx)
                                    next_tcPr = next_row_cell_obj._tc.tcPr
                                    if next_tcPr is not None and next_tcPr.vMerge is not None and next_tcPr.vMerge.get(qn('w:val')) != 'restart':
                                        rowspan += 1
                                    else:
                                        break # End of this vertical span
                                except IndexError:
                                    break # Row or cell doesn't exist, end of span
                        
                        # Mark the grid cells occupied by this primary cell and its spans
                        for r_offset in range(rowspan):
                            for c_offset in range(colspan):
                                if (r_idx + r_offset) < len(table_object.rows) and (c_idx + c_offset) < actual_cols:
                                    if grid_cell_occupier[r_idx + r_offset][c_idx + c_offset] is None:
                                        grid_cell_occupier[r_idx + r_offset][c_idx + c_offset] = (r_idx, c_idx)
                                    elif grid_cell_occupier[r_idx + r_offset][c_idx + c_offset] != (r_idx, c_idx):
                                        logger.warning(f"  Overlap detected at grid ({r_idx + r_offset},{c_idx + c_offset}) for table {doc_table_index}. Original occupier: {grid_cell_occupier[r_idx + r_offset][c_idx + c_offset]}, new primary: ({r_idx},{c_idx}).")
                                else:
                                    logger.warning(f"  Span of cell ({r_idx},{c_idx}) in table {doc_table_index} goes out of bounds at ({r_idx + r_offset},{c_idx + c_offset}).")

                        # Extract combined content from this primary cell
                        combined_cell_text = []
                        combined_cell_runs = []
                        first_para_style = None
                        first_para_alignment = None
                        first_para_page_break_before = False

                        if not current_cell_obj.paragraphs:
                            logger.debug(f"  Primary cell ({r_idx},{c_idx}) in table {doc_table_index} (TableMetaID: {current_block_id_for_table_meta}) has no paragraphs. Creating empty block.")
                            # Still create a block for it to represent structure and allow editing empty cells
                        else:
                            for cp_idx, cell_para_element in enumerate(current_cell_obj.paragraphs):
                                combined_cell_text.append(cell_para_element.text)
                                if cp_idx == 0:
                                    first_para_style = cell_para_element.style.name if cell_para_element.style else "Normal"
                                    # --- NEW: Check for run-level page breaks (w:br w:type="page") in first paragraph of cell ---
                                    cell_pPr = cell_para_element._element.pPr
                                    cell_pbf_value = cell_pPr.pageBreakBefore is not None if cell_pPr is not None else False
                                    first_para_page_break_before = cell_pbf_value

                                    if cell_para_element.alignment is not None:
                                        alignment_map = {
                                            WD_PARAGRAPH_ALIGNMENT.LEFT: "LEFT", WD_PARAGRAPH_ALIGNMENT.CENTER: "CENTER",
                                            WD_PARAGRAPH_ALIGNMENT.RIGHT: "RIGHT", WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "JUSTIFY",
                                            WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: "DISTRIBUTE", WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY: "THAI_JUSTIFY"
                                        }
                                        first_para_alignment = alignment_map.get(cell_para_element.alignment)
                                
                                for run_element in cell_para_element.runs:
                                    font = run_element.font
                                    combined_cell_runs.append({
                                        "text": run_element.text,
                                        "bold": run_element.bold if run_element.bold is not None else False,
                                        "italic": run_element.italic if run_element.italic is not None else False,
                                        "underline": run_element.underline if run_element.underline is not None else False,
                                        "font_name": font.name,
                                        "font_size_pt": font.size.pt if font.size and hasattr(font.size, 'pt') else None,
                                        "font_color_rgb": str(font.color.rgb) if font.color and font.color.rgb else None
                                    })
                        
                        final_combined_text = "\n".join(combined_cell_text)
                        
                        cell_content_block: Dict[str, Any] = {
                            "id": f"tc_block_{block_id_counter}",
                            "doc_table_index": doc_table_index, 
                            "row_index": r_idx, # This is the primary cell's row index
                            "col_index": c_idx, # This is the primary cell's col index
                            "row_span": rowspan, # Store calculated rowspan
                            "col_span": colspan, # Store calculated colspan
                            "overall_block_index": block_id_counter,
                            "type": "table_cell_combined",
                            "text": final_combined_text,
                            "style_name": first_para_style, 
                            "alignment": first_para_alignment,
                            "page_break_before": first_para_page_break_before, # Ensure this uses the debugged value
                            "runs": combined_cell_runs
                        }
                        content_blocks.append(cell_content_block)
                        logger.debug(f"  Created tc_block_{block_id_counter} for primary cell ({r_idx},{c_idx}), Table {doc_table_index} (MetaID: {current_block_id_for_table_meta}), Colspan: {colspan}, Rowspan: {rowspan}")
                        block_id_counter += 1
            else:
                logger.debug(f"Skipping unexpected element type in document body: {type(child_element)}")

        logger.info(f"Extraction complete. Found {len(content_blocks)} total structured blocks by iterating body elements.")
        return content_blocks

    def _apply_formatting_to_paragraph(self, para_to_edit: Paragraph, new_text: str, 
                                   original_runs_info: List[Dict[str, Any]],
                                   original_para_style_name: Optional[str] = None,
                                      original_para_alignment: Optional[str] = None,
                                      original_page_break_before: Optional[bool] = None):
        """Helper function to clear and apply formatting to a paragraph object."""
        
        # Clear existing runs
        while para_to_edit.runs:
            run_to_remove = para_to_edit.runs[0]
            para_to_edit._p.remove(run_to_remove._r)

        original_full_text = "".join([r_info.get("text", "") for r_info in original_runs_info])

        # Apply new text and formatting
        if new_text == original_full_text and original_runs_info:
            logger.debug("Text unchanged, reapplying original run formatting.")
            for r_info in original_runs_info:
                run_text = r_info.get("text", "")
                added_run = para_to_edit.add_run(run_text)
                added_run.bold = r_info.get('bold', False)
                added_run.italic = r_info.get('italic', False)
                added_run.underline = r_info.get('underline', False)
                if r_info.get('font_name'):
                    added_run.font.name = r_info['font_name']
                    added_run._element.rPr.rFonts.set(qn('w:eastAsia'), r_info['font_name'])
                if r_info.get('font_size_pt'):
                    added_run.font.size = Pt(r_info['font_size_pt'])
                if r_info.get('font_color_rgb'):
                    try:
                        color_str = r_info['font_color_rgb']
                        if color_str.startswith("RGBColor("): 
                            parts = color_str.replace("RGBColor(", "").replace(")", "").split(',')
                            added_run.font.color.rgb = RGBColor(int(parts[0].strip(),16), int(parts[1].strip(),16), int(parts[2].strip(),16))
                        elif len(color_str) == 6 and all(c in '0123456789abcdefABCDEF' for c in color_str):
                             added_run.font.color.rgb = RGBColor.from_string(color_str)
                        elif color_str: logger.warning(f"Unrecognized RGB color string format '{color_str}' for run, skipping.")
                    except ValueError as ve: logger.warning(f"Invalid RGB color string '{r_info.get('font_color_rgb')}': {ve}")
        else:
            logger.debug("Text changed. Applying new text with formatting from first original run (if any).")
            added_run = para_to_edit.add_run(new_text)
            if original_runs_info:
                first_run_info = original_runs_info[0]
                added_run.bold = first_run_info.get('bold', False)
                added_run.italic = first_run_info.get('italic', False)
                added_run.underline = first_run_info.get('underline', False)
                if first_run_info.get('font_name'):
                    added_run.font.name = first_run_info['font_name']
                    added_run._element.rPr.rFonts.set(qn('w:eastAsia'), first_run_info['font_name'])
                if first_run_info.get('font_size_pt'):
                    added_run.font.size = Pt(first_run_info['font_size_pt'])
                if first_run_info.get('font_color_rgb'):
                    try:
                        color_str = first_run_info['font_color_rgb']
                        if color_str.startswith("RGBColor("):
                            parts = color_str.replace("RGBColor(", "").replace(")", "").split(',')
                            added_run.font.color.rgb = RGBColor(int(parts[0].strip(),16), int(parts[1].strip(),16), int(parts[2].strip(),16))
                        elif len(color_str) == 6 and all(c in '0123456789abcdefABCDEF' for c in color_str):
                             added_run.font.color.rgb = RGBColor.from_string(color_str)
                        elif color_str: logger.warning(f"Unrecognized RGB color string format '{color_str}' for first run, skipping.")
                    except ValueError as ve: logger.warning(f"Invalid RGB color string '{first_run_info.get('font_color_rgb')}': {ve}")

        # Re-apply paragraph-level style
        if original_para_style_name and self.current_document:
            try:
                available_style_names = [s.name for s in self.current_document.styles]
                if original_para_style_name in available_style_names:
                    if para_to_edit.style.name != original_para_style_name:
                        para_to_edit.style = self.current_document.styles[original_para_style_name]
                        logger.debug(f"Applied style '{original_para_style_name}'.")
                else:
                    logger.warning(f"Style '{original_para_style_name}' not found. Current: '{para_to_edit.style.name}'.")
            except Exception as e_style:
                logger.error(f"Error applying paragraph style '{original_para_style_name}': {e_style}")

        # Re-apply paragraph-level alignment
        if original_para_alignment and self.current_document:
            alignment_map_reverse = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT, "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT, "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                "DISTRIBUTE": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE, "THAI_JUSTIFY": WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY
            }
            align_val = alignment_map_reverse.get(original_para_alignment.upper())
            if align_val is not None:
                 para_to_edit.alignment = align_val
            else: logger.warning(f"Unknown alignment value: {original_para_alignment}")

        # Re-apply page_break_before
        if original_page_break_before is True:
            para_to_edit.paragraph_format.page_break_before = True
            logger.debug(f"Applied page_break_before=True.")
        elif original_page_break_before is False: # Explicitly set to false if it was false
            para_to_edit.paragraph_format.page_break_before = False

    def edit_block_content_internal(self, 
                                   new_text: str, 
                                   original_runs_info: List[Dict[str, Any]],
                                   # --- Identification Args (mutually exclusive) ---
                                   doc_paragraph_index: Optional[int] = None, 
                                   doc_table_index: Optional[int] = None,
                                   row_index: Optional[int] = None,
                                   col_index: Optional[int] = None,
                                   # --- Formatting Args ---
                                   original_para_style_name: Optional[str] = None,
                                   original_para_alignment: Optional[str] = None,
                                   original_page_break_before: Optional[bool] = None) -> None:
        """
        Internal logic to edit a specific paragraph, either top-level or inside a table cell.
        Finds the target paragraph using indices and applies text/formatting.
        """
        if not self.current_document:
            logger.warning("edit_block_content_internal: No active document.")
            raise ValueError("No active document to process.")

        para_to_edit: Optional[Paragraph] = None
        identifier_log = ""

        # --- Find the Paragraph to Edit ---
        if doc_paragraph_index is not None:
            # Editing a top-level paragraph
            identifier_log = f"top-level paragraph index {doc_paragraph_index}"
            if 0 <= doc_paragraph_index < len(self.current_document.paragraphs):
                para_to_edit = self.current_document.paragraphs[doc_paragraph_index]
                logger.debug(f"Editing content for {identifier_log}")
                self._apply_formatting_to_paragraph(
                    para_to_edit, new_text, original_runs_info,
                    original_para_style_name, original_para_alignment,
                    original_page_break_before
                )
            else:
                logger.error(f"edit_block_content_internal: Top-level paragraph index {doc_paragraph_index} out of range.")
                raise IndexError(f"Top-level paragraph index {doc_paragraph_index} out of range.")
        
        elif doc_table_index is not None and row_index is not None and col_index is not None:
            # Editing a combined table cell content
            # cell_paragraph_index is no longer used for combined cells
            identifier_log = f"table {doc_table_index}, cell ({row_index},{col_index})"
            if 0 <= doc_table_index < len(self.current_document.tables):
                table = self.current_document.tables[doc_table_index]
                if 0 <= row_index < len(table.rows):
                    if 0 <= col_index < len(table.columns): 
                        cell_to_edit = table.cell(row_index, col_index)
                        logger.debug(f"Editing content for {identifier_log}")
                        
                        # Clear existing paragraphs in the cell
                        # A cell's content is primarily its paragraphs. To replace cell content,
                        # we clear existing paragraphs and add one new one with the new_text.
                        # Accessing private _element and _tc to remove paragraph elements directly.
                        for p_element in cell_to_edit._tc.xpath('./w:p'):
                            cell_to_edit._tc.remove(p_element)
                        
                        # Add a new paragraph with the new_text and apply formatting
                        # The formatting (runs, style, alignment) will be applied to this new single paragraph.
                        # If new_text itself contains newlines, they will be preserved in the new paragraph.
                        # If the original cell had multiple paragraphs, they are now merged into one.
                        new_para_in_cell = cell_to_edit.add_paragraph() # Add new, empty paragraph
                        self._apply_formatting_to_paragraph(
                            new_para_in_cell, new_text, original_runs_info,
                            original_para_style_name, original_para_alignment,
                            original_page_break_before
                        )
                    else:
                        logger.error(f"edit_block_content_internal: Column index {col_index} out of range for table {doc_table_index}.")
                        raise IndexError(f"Column index {col_index} out of range.")
                else:
                    logger.error(f"edit_block_content_internal: Row index {row_index} out of range for table {doc_table_index}.")
                    raise IndexError(f"Row index {row_index} out of range.")
            else:
                logger.error(f"edit_block_content_internal: Table index {doc_table_index} out of range.")
                raise IndexError(f"Table index {doc_table_index} out of range.")
        else:
            # Invalid combination of arguments
            logger.error("edit_block_content_internal: Invalid arguments. Must provide 'doc_paragraph_index' OR ('doc_table_index', 'row_index', 'col_index').")
            raise ValueError("Invalid arguments for identifying content to edit.")

        # --- Apply Formatting to the Found Paragraph ---
        # This section is removed as the logic is now split and called within the if/elif blocks above
        # if para_to_edit is not None: ... 

# Create global processor instance
processor = DocxProcessor()

@asynccontextmanager
async def server_lifespan(server: FastMCP) -> AsyncIterator[Dict[str, Any]]:
    """Manage server lifecycle"""
    global processor
    try:
        logger.info("DocxProcessor MCP server starting...")
        yield {"processor": processor}
    finally:
        logger.info("DocxProcessor MCP server shutting down...")
        processor.save_state()

# Create MCP server
mcp = FastMCP(
    name="DocxProcessor",
    instructions="Word document processing service, providing functions to create, edit, and query documents",
    lifespan=server_lifespan,
    port=8001
)

# --- Direct HTTP Endpoint Implementations ---

async def http_open_document(request: StarletteRequest) -> JSONResponse:
    try:
        data = await request.json()
        file_path = data.get("file_path")
        if not file_path:
            return JSONResponse({"status": "error", "message": "file_path is required"}, status_code=400)
        
        if not os.path.exists(file_path):
            logger.warning(f"Direct HTTP: File does not exist: {file_path}")
            return JSONResponse({"status": "error", "message": f"File does not exist: {file_path}"}, status_code=404)
        
        processor.current_document = Document(file_path)
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        processor._save_current_document_path_state()
        
        logger.info(f"Direct HTTP: Document opened successfully: {file_path}")
        return JSONResponse({"status": "success", "message": f"Document opened successfully: {file_path}"})
    except Exception as e:
        error_msg = f"Direct HTTP: Failed to open document: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return JSONResponse({"status": "error", "message": error_msg}, status_code=500)

async def http_get_structured_content(request: StarletteRequest) -> JSONResponse:
    try:
        if not processor.current_document:
            logger.warning("Direct HTTP: get_structured_document_content: No document is open")
            return JSONResponse({"status": "error", "message": "No document is open"}, status_code=400)
        
        structured_content = processor.get_structured_document_content_internal()
        logger.info(f"Direct HTTP: Successfully extracted {len(structured_content)} blocks.")
        return JSONResponse({"status": "success", "content": structured_content})
    except Exception as e:
        error_msg = f"Direct HTTP: Failed to get structured document content: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return JSONResponse({"status": "error", "message": error_msg, "trace": traceback.format_exc() if logger.level == logging.DEBUG else None}, status_code=500)

async def http_edit_block_content(request: StarletteRequest) -> JSONResponse:
    """HTTP endpoint to edit content of a paragraph (top-level or in table)."""
    try:
        data = await request.json()
        new_text = data.get("new_text")
        original_runs_info = data.get("original_runs_info", [])
        original_para_style_name = data.get("original_para_style_name")
        original_para_alignment = data.get("original_para_alignment")
        original_page_break_before = data.get("original_page_break_before", False)

        # Identification parameters (mutually exclusive)
        doc_paragraph_index = data.get("doc_paragraph_index")
        doc_table_index = data.get("doc_table_index")
        row_index = data.get("row_index")
        col_index = data.get("col_index")

        # Basic validation
        if new_text is None:
            return JSONResponse({"status": "error", "message": "new_text is required."}, status_code=400)
        
        is_paragraph_edit = doc_paragraph_index is not None
        is_table_cell_edit = (doc_table_index is not None and 
                              row_index is not None and 
                              col_index is not None)

        if not is_paragraph_edit and not is_table_cell_edit:
            return JSONResponse({
                "status": "error", 
                "message": "Invalid identification parameters. Provide either 'doc_paragraph_index' or all of ('doc_table_index', 'row_index', 'col_index')."
            }, status_code=400)
        
        if is_paragraph_edit and is_table_cell_edit:
             return JSONResponse({"status": "error", "message": "Cannot provide both paragraph and table cell identifiers."}, status_code=400)

        if not processor.current_document:
            logger.warning(f"Direct HTTP: edit_block_content: No document is open")
            return JSONResponse({"status": "error", "message": "No document is open"}, status_code=400)

        # Call the internal logic with appropriate arguments
        processor.edit_block_content_internal(
            new_text=new_text, 
            original_runs_info=original_runs_info,
            doc_paragraph_index=doc_paragraph_index, 
            doc_table_index=doc_table_index,
            row_index=row_index,
            col_index=col_index,
            original_para_style_name=original_para_style_name,
            original_para_alignment=original_para_alignment,
            original_page_break_before=original_page_break_before
        )
        
        # Determine identifier for log message
        if is_paragraph_edit:
            edit_location = f"paragraph index {doc_paragraph_index}"
        elif is_table_cell_edit:
            # cell_paragraph_index no longer used for combined cells
            edit_location = f"table index {doc_table_index}, cell ({row_index},{col_index})"
        else:
            edit_location = "unknown location" # Should not happen due to earlier checks
            
        logger.info(f"Direct HTTP: Successfully edited content for {edit_location}.")
        return JSONResponse({"status": "success", "message": f"Content at {edit_location} updated successfully."})

    except (IndexError, ValueError) as e_val_idx:
        error_msg = f"Direct HTTP: Failed to edit content due to invalid index or value: {str(e_val_idx)}"
        logger.error(error_msg, exc_info=False)
        return JSONResponse({"status": "error", "message": error_msg}, status_code=400)
    except Exception as e:
        error_msg = f"Direct HTTP: Unexpected error editing content: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return JSONResponse({"status": "error", "message": error_msg, "trace": traceback.format_exc() if logger.level == logging.DEBUG else None}, status_code=500)

async def http_save_as_document(request: StarletteRequest) -> JSONResponse:
    try:
        data = await request.json()
        new_file_path = data.get("new_file_path")
        if not new_file_path:
            return JSONResponse({"status": "error", "message": "new_file_path is required"}, status_code=400)

        if not processor.current_document:
            logger.warning("Direct HTTP: No document open to save as.")
            return JSONResponse({"status": "error", "message": "No document is open"}, status_code=400)
        
        processor.current_document.save(new_file_path)
        
        processor.current_file_path = new_file_path # Update current path
        processor.documents[new_file_path] = processor.current_document # Update documents dict
        processor._save_current_document_path_state() # Add this
        
        logger.info(f"Direct HTTP: Document saved as: {new_file_path}")
        return JSONResponse({"status": "success", "message": f"Document saved as: {new_file_path}", "file_path": new_file_path})
    except Exception as e:
        error_msg = f"Direct HTTP: Failed to save document as: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return JSONResponse({"status": "error", "message": error_msg}, status_code=500)

# This definition needs to be moved before the __main__ block
direct_http_app = StarletteApp(routes=[
    HttpRoute("/api/open_document", http_open_document, methods=["POST"]),
    HttpRoute("/api/get_structured_content", http_get_structured_content, methods=["GET"]),
    HttpRoute("/api/edit_block_content", http_edit_block_content, methods=["POST"]),
    HttpRoute("/api/save_as_document", http_save_as_document, methods=["POST"]),
])

# Uvicorn logging configuration to output to stderr
UVICORN_LOGGING_CONFIG = {
    "version": 1,
    "disable_existing_loggers": False,
    "formatters": {
        "default": {
            "()": "uvicorn.logging.DefaultFormatter",
            "fmt": "%(levelprefix)s %(message)s",
            "use_colors": None,
        },
        "access": {
            "()": "uvicorn.logging.AccessFormatter",
            # Exclude client_addr and request_line for cleaner MCP logs, keep status_code
            "fmt": '''%(levelprefix)s %(status_code)s''',
        },
    },
    "handlers": {
        "default": {
            "formatter": "default",
            "class": "logging.StreamHandler",
            "stream": "ext://sys.stderr",
        },
        "access": {
            "formatter": "access",
            "class": "logging.StreamHandler",
            "stream": "ext://sys.stderr",
        },
    },
    "loggers": {
        "uvicorn": {"handlers": ["default"], "level": "INFO", "propagate": False},
        "uvicorn.error": {"handlers": ["default"], "level": "INFO", "propagate": False}, # Ensure uvicorn errors also use the default handler (stderr)
        "uvicorn.access": {"handlers": ["access"], "level": "WARNING", "propagate": False}, # Reduce verbosity of access logs or set level to ERROR
    },
}

if __name__ == "__main__":
    logger.info("Attempting to start DocxProcessor server with combined MCP-SSE and Direct HTTP interfaces.")
    if os.path.exists(CURRENT_DOC_FILE):
        try:
            os.remove(CURRENT_DOC_FILE)
            logger.info(f"Removed existing state file '{CURRENT_DOC_FILE}' for clean startup.")
        except Exception as e:
            logger.error(f"Failed to remove existing state file '{CURRENT_DOC_FILE}': {e}")
    
    try:
        # Main application to run with Uvicorn
        main_app = StarletteApp()
        
        # Mount the FastMCP SSE application (e.g., for other MCP compliant clients)
        # FastMCP's sse_app() uses its own settings for host/port if not overridden by uvicorn.run
        # It defines routes like /sse and /messages/
        main_app.mount("/mcp", app=mcp.sse_app()) 
        
        # Mount the direct HTTP API
        main_app.mount("/", app=direct_http_app) # Mount direct API at root or another path like /direct_api

        uvicorn.run(main_app, host="0.0.0.0", port=8001, log_config=UVICORN_LOGGING_CONFIG)
        
        logger.info("Uvicorn server with combined interfaces has finished or exited.") 
    except Exception as e:
        logger.error(f"Exception during Uvicorn server startup: {e}", exc_info=True)